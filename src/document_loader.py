"""
Document loader and chunking system for Agentic RAG
"""
from pathlib import Path
from typing import List, Optional, Dict
import PyPDF2
import docx
import json
from dataclasses import dataclass, asdict
from datetime import datetime
import hashlib


@dataclass
class Document:
    """Represents a document chunk"""
    id: str
    content: str
    source: str
    chunk_index: int
    total_chunks: int
    metadata: Dict = None

    def to_dict(self) -> Dict:
        return asdict(self)


class DocumentLoader:
    """Loads and chunks documents from various formats"""

    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def load_pdf(self, file_path: str) -> str:
        """Load text from PDF"""
        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text += page.extract_text()
        return text

    def load_docx(self, file_path: str) -> str:
        """Load text from DOCX"""
        doc = docx.Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text

    def load_txt(self, file_path: str) -> str:
        """Load text from TXT"""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()

    def load_md(self, file_path: str) -> str:
        """Load text from Markdown"""
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read()

    def load_document(self, file_path: str) -> Optional[str]:
        """Load document based on file type"""
        file_path = Path(file_path)
        
        if file_path.suffix.lower() == ".pdf":
            return self.load_pdf(file_path)
        elif file_path.suffix.lower() == ".docx":
            return self.load_docx(file_path)
        elif file_path.suffix.lower() in [".txt", ".md"]:
            return self.load_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_path.suffix}")

    def chunk_text(self, text: str, source: str) -> List[Document]:
        """Split text into overlapping chunks"""
        text = text.strip()
        if not text:
            return []

        step = max(1, self.chunk_size - self.chunk_overlap)
        start_positions = list(range(0, len(text), step))
        total_chunks = len(start_positions)
        chunks = []

        for i, start_idx in enumerate(start_positions):
            end_idx = min(start_idx + self.chunk_size, len(text))
            chunk_content = text[start_idx:end_idx]

            chunk_id = f"{source}_{i}_{hashlib.md5(chunk_content.encode()).hexdigest()[:8]}"
            doc = Document(
                id=chunk_id,
                content=chunk_content,
                source=source,
                chunk_index=i,
                total_chunks=total_chunks,
                metadata={
                    "timestamp": datetime.now().isoformat(),
                    "chunk_size": len(chunk_content),
                    "start_char": start_idx,
                    "end_char": end_idx
                }
            )
            chunks.append(doc)

        return chunks

    def process_documents(self, document_dir: str) -> List[Document]:
        """Process all documents in a directory"""
        all_chunks = []
        doc_dir = Path(document_dir)
        
        # Supported file extensions
        supported_extensions = [".pdf", ".docx", ".txt", ".md"]
        
        for file_path in doc_dir.rglob("*"):
            if file_path.suffix.lower() in supported_extensions:
                try:
                    print(f"Processing: {file_path}")
                    text = self.load_document(str(file_path))
                    source = file_path.name
                    chunks = self.chunk_text(text, source)
                    all_chunks.extend(chunks)
                    print(f"  - Created {len(chunks)} chunks")
                except Exception as e:
                    print(f"Error processing {file_path}: {str(e)}")
        
        return all_chunks

    def save_chunks(self, chunks: List[Document], output_path: str):
        """Save chunks to JSON for reference"""
        output_file = Path(output_path) / "chunks.jsonl"
        with open(output_file, "w", encoding="utf-8") as f:
            for chunk in chunks:
                f.write(json.dumps(chunk.to_dict()) + "\n")
        print(f"Saved {len(chunks)} chunks to {output_file}")
